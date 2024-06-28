import re

import polars as pl
from docx import Document

DF_PATH = "etl/data/gold/olist_product_reviews_dataset.parquet"
REVIEW_PATH = "ai/data/reviews.txt"
REPORT_PATH = "ai/data/report.docx"


class ReviewsExtractor:
    def __init__(self):
        pass

    def read_table(self, product_id):
        df = pl.read_parquet(DF_PATH)
        df = df.filter(pl.col("product_id") == f"{product_id}")
        return df

    def reviews_to_txt(self, df):
        product_reviews = df["review_comment_message"].to_list()

        with open(REVIEW_PATH, "w") as file:
            for line in product_reviews:
                file.write(line + "\n")

    def load_txt(self):
        with open(REVIEW_PATH, "r") as f:
            reviews = f.read()
            return reviews

    def extract_section(self, section_title, ai_response):
        pattern = r"\*\*{}:\*\*\s*([\s\S]*?)(?=\*\*|$)".format(section_title)
        match = re.search(pattern, ai_response)
        if match:
            return match.group(1).strip()

    def build_report(self, ai_response):
        doc = Document()

        doc.add_heading("Avaliações de produto", level=1)

        # Add Pontos Positivos
        pontos_positivos = self.extract_section("Pontos Positivos", ai_response)
        if pontos_positivos:
            doc.add_heading("Pontos Positivos", level=2)
            for item in pontos_positivos.split("\n"):
                if item.strip().startswith("* "):
                    doc.add_paragraph(item.strip("* "), style="List Bullet")

        # Add Pontos Negativos
        pontos_negativos = self.extract_section("Pontos Negativos", ai_response)
        if pontos_negativos:
            doc.add_heading("Pontos Negativos", level=2)
            for item in pontos_negativos.split("\n"):
                if item.strip().startswith("* "):
                    doc.add_paragraph(item.strip("* "), style="List Bullet")

        # Add Recomendações
        recomendacoes = self.extract_section("Recomendações", ai_response)
        if recomendacoes:
            doc.add_heading("Recomendações", level=2)
            for item in recomendacoes.split("\n"):
                if item.strip().startswith("* "):
                    doc.add_paragraph(item.strip("* "), style="List Bullet")
        """
        # Add extracted reviews
        reviews = self.load_txt()
        doc.add_heading('Customer Reviews', level=2)
        for review in reviews.split('\n'):
            doc.add_paragraph(review)
        """
        doc.save(REPORT_PATH)
